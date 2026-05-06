from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

BASE = Path(r"D:\pak if\RPL XI\BLOK VII")
SRC = BASE / "LAPORAN ANALISIS DAN SISTEM PERANCANGAN SISTEM INFORMASI ABSENSI PKL.docx"
BACKUP = BASE / "LAPORAN ANALISIS DAN SISTEM PERANCANGAN SISTEM INFORMASI ABSENSI PKL - BACKUP SEBELUM REVISI.docx"
OUT = BASE / "LAPORAN ANALISIS DAN PERANCANGAN SISTEM INFORMASI ABSENSI PKL - REVISI.docx"
DESIGN = BASE / "PERANCANGAN ABSENSI PKL"
SCREENS = BASE / "FOTO HALAMAN APLIKASI ABSENSI PKL"

if not BACKUP.exists():
    shutil.copy2(SRC, BACKUP)

doc = Document(SRC)
body = doc._body._element
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

for name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Caption"]:
    if name in doc.styles:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12 if name == "Normal" else 11)

TERMS = ["black box testing", "dashboard", "database", "flowchart", "online", "offline", "login", "input", "output", "role", "website", "web", "backup", "restore", "export", "import", "user"]
PAT = re.compile("(" + "|".join(re.escape(x) for x in sorted(TERMS, key=len, reverse=True)) + ")", re.I)


def runs(p, text, italic=True):
    if not italic:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        return
    pos = 0
    for m in PAT.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            r.font.name = "Times New Roman"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r = p.add_run(m.group(0))
        r.italic = True
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def p(text="", style="Normal", align=None, italic=True):
    para = doc.add_paragraph(style=style)
    runs(para, text, italic)
    para.paragraph_format.line_spacing = 1.5 if style == "Normal" else None
    para.paragraph_format.space_after = Pt(6)
    para.alignment = align if align is not None else (WD_ALIGN_PARAGRAPH.JUSTIFY if style == "Normal" else None)
    return para


def h(text, level=1):
    return p(text, f"Heading {level}", italic=False)


def page():
    doc.add_page_break()


def new_section():
    return doc.add_section(WD_SECTION.NEW_PAGE)


def clear_range(rng):
    for para in rng.paragraphs:
        para.text = ""


def add_page_field(paragraph):
    paragraph.text = ""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def setup_page_numbers():
    for index, section in enumerate(doc.sections):
        section.different_first_page_header_footer = True
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False

        clear_range(section.header)
        clear_range(section.footer)
        clear_range(section.first_page_header)
        clear_range(section.first_page_footer)

        if index == 0:
            continue

        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_field(hp)

        fp = section.first_page_footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_field(fp)


def simple_field(title, instr):
    h(title, 1)
    para = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Daftar akan diperbarui otomatis oleh Microsoft Word."
    r.append(t)
    fld.append(r)
    para._p.append(fld)
    new_section()


def img(path, caption, text, width=5.8):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Inches(width))
    p(caption, "Caption", WD_ALIGN_PARAGRAPH.CENTER, italic=False)
    p(text)


def numbered(items):
    for i, x in enumerate(items, 1):
        p(f"{i}. {x}")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        t.rows[0].cells[i].text = head
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
            r.font.name = "Times New Roman"
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                for r in para.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(9)
    doc.add_paragraph()


def explain(caption, kind):
    return (
        f"{caption} menunjukkan bagian {kind} pada Sistem Informasi Absensi PKL yang digunakan untuk mendukung proses pencatatan dan monitoring kehadiran siswa. "
        "Bagian ini menampilkan alur, data, atau fitur yang berkaitan langsung dengan kebutuhan pengguna dalam menjalankan kegiatan PKL. "
        "Informasi yang ditampilkan membantu pengguna memahami fungsi sistem secara lebih jelas dan terstruktur. "
        "Dengan adanya bagian ini, proses administrasi absensi PKL dapat dilakukan lebih rapi, cepat, dan mudah dievaluasi oleh pihak sekolah."
    )


# Cover
for _ in range(2):
    p("", align=WD_ALIGN_PARAGRAPH.CENTER)
r = p("LAPORAN ANALISIS DAN PERANCANGAN SISTEM INFORMASI ABSENSI PKL", align=WD_ALIGN_PARAGRAPH.CENTER, italic=False)
r.runs[0].bold = True
for _ in range(5):
    p("", align=WD_ALIGN_PARAGRAPH.CENTER)
p("Disusun oleh :", align=WD_ALIGN_PARAGRAPH.CENTER, italic=False)
p("Rachel Keithlin", align=WD_ALIGN_PARAGRAPH.CENTER, italic=False)
p("24161025", align=WD_ALIGN_PARAGRAPH.CENTER, italic=False)
for _ in range(8):
    p("", align=WD_ALIGN_PARAGRAPH.CENTER)
for x in ["REKAYASA PERANGKAT LUNAK", "SMK PERMATA HARAPAN", "BATAM", "2025"]:
    q = p(x, align=WD_ALIGN_PARAGRAPH.CENTER, italic=False)
    q.runs[0].bold = True
new_section()

h("KATA PENGANTAR", 1)
p("Puji syukur saya panjatkan ke hadirat Tuhan Yang Maha Esa karena atas rahmat dan karunia-Nya laporan analisis dan perancangan Sistem Informasi Absensi PKL ini dapat diselesaikan dengan baik. Laporan ini disusun sebagai dokumentasi proses analisis kebutuhan, perancangan sistem, dan hasil implementasi aplikasi absensi PKL. Penyusunan laporan ini juga menjadi bagian dari pembelajaran pada kompetensi Rekayasa Perangkat Lunak. Dengan adanya laporan ini, pembaca diharapkan memperoleh gambaran jelas mengenai kebutuhan dan rancangan sistem informasi absensi PKL.")
p("Saya mengucapkan terima kasih kepada pihak sekolah, guru pembimbing, dan semua pihak yang telah memberikan arahan selama proses penyusunan laporan ini. Bantuan berupa masukan dan bimbingan sangat membantu dalam memperbaiki isi laporan. Saya juga berterima kasih kepada teman-teman yang telah memberikan dukungan selama proses pengumpulan data dan penyusunan rancangan. Setiap masukan yang diberikan menjadi bahan pertimbangan agar laporan ini lebih sesuai dengan topik Sistem Informasi Absensi PKL.")
p("Saya menyadari bahwa laporan ini masih memiliki kekurangan, baik dari segi penulisan maupun kelengkapan pembahasan. Oleh karena itu, kritik dan saran yang membangun sangat diharapkan agar laporan ini dapat menjadi lebih baik pada masa mendatang. Laporan ini disusun dengan menggunakan bahasa formal dan mengikuti format laporan yang telah tersedia sebelumnya. Semoga laporan ini dapat bermanfaat bagi pembaca dalam memahami perancangan sistem informasi untuk membantu proses absensi siswa PKL.")
p("Batam, 12 Agustus 2025", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=False)
p("Rachel Keithlin", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=False)
new_section()
simple_field("DAFTAR ISI", r'TOC \o "1-3" \h \z \u')
simple_field("DAFTAR GAMBAR", r'TOC \h \z \t "Caption,1"')

# BAB I
h("BAB I\nPENDAHULUAN", 1)
h("1.1 Latar Belakang", 2)
latars = [
    "Praktik Kerja Lapangan merupakan kegiatan penting bagi siswa SMK karena memberikan pengalaman langsung di dunia kerja. Melalui kegiatan ini, siswa menerapkan kompetensi yang dipelajari di sekolah pada situasi kerja nyata. Kegiatan PKL juga melatih kedisiplinan, tanggung jawab, komunikasi, dan kemampuan menyelesaikan tugas sesuai bidang keahlian. Oleh karena itu, pelaksanaan PKL perlu didukung administrasi yang tertib agar sekolah dapat memantau perkembangan siswa dengan baik.",
    "Salah satu administrasi penting dalam kegiatan PKL adalah pencatatan kehadiran siswa. Kehadiran menjadi indikator kedisiplinan dan bukti bahwa siswa mengikuti kegiatan sesuai jadwal. Data kehadiran diperlukan oleh admin, kesiswaan, guru pembimbing, dan sekolah sebagai dasar evaluasi pelaksanaan PKL. Apabila pencatatan tidak berjalan baik, sekolah akan kesulitan memperoleh informasi yang akurat mengenai aktivitas siswa.",
    "Pencatatan absensi PKL yang masih manual memiliki banyak kendala dalam pelaksanaannya. Data dapat dicatat pada buku, lembar presensi, atau laporan terpisah dari tempat praktik. Cara tersebut berisiko menyebabkan data hilang, rusak, terlambat dikumpulkan, atau tidak sesuai dengan kondisi sebenarnya. Proses rekapitulasi juga membutuhkan waktu lama karena petugas harus memeriksa data dari banyak sumber.",
    "Kesulitan monitoring menjadi masalah lain ketika siswa ditempatkan pada banyak lokasi PKL. Pihak sekolah tidak selalu dapat memeriksa kehadiran siswa secara langsung setiap hari. Guru pembimbing membutuhkan informasi yang cepat untuk mengetahui siswa yang hadir, terlambat, izin, sakit, atau tidak hadir. Jika informasi terlambat diterima, tindak lanjut terhadap masalah kehadiran juga dapat terlambat dilakukan.",
    "Data absensi yang tidak akurat dapat memengaruhi proses evaluasi kegiatan PKL. Kesalahan pencatatan dapat menyebabkan siswa yang hadir dianggap tidak hadir atau sebaliknya. Kondisi tersebut dapat menimbulkan perbedaan data antara siswa, pembimbing, tempat PKL, dan sekolah. Karena itu, sistem pencatatan yang terstruktur diperlukan agar data dapat disimpan, diperiksa, dan digunakan kembali secara tepat.",
    "Kebutuhan laporan absensi juga menjadi alasan penting dibuatnya sistem informasi. Laporan kehadiran diperlukan dalam bentuk harian, mingguan, atau periode tertentu sesuai kebutuhan sekolah. Jika data masih tersebar pada dokumen manual, pembuatan laporan akan memerlukan waktu panjang dan rentan salah hitung. Sistem terkomputerisasi dapat mempercepat proses tersebut karena data tersimpan dalam database yang dapat diolah secara otomatis.",
    "Sistem informasi dapat membantu mengatasi kendala absensi manual karena data dikelola secara terpusat. Data siswa, lokasi PKL, absensi masuk, absensi pulang, status kehadiran, dan laporan dapat saling terhubung. Pengguna dapat mengakses informasi sesuai role dan hak akses masing-masing. Dengan cara ini, proses pencatatan dan monitoring tidak lagi bergantung sepenuhnya pada dokumen manual.",
    "Sistem Informasi Absensi PKL dirancang untuk membantu siswa melakukan pencatatan kehadiran secara praktis. Siswa dapat melakukan absensi masuk dan pulang melalui halaman yang telah disediakan. Admin atau kesiswaan dapat memantau data yang masuk, memeriksa riwayat, serta mengelola data pendukung seperti siswa dan lokasi PKL. Guru pembimbing juga dapat menggunakan informasi tersebut untuk mendukung pengawasan selama siswa berada di tempat praktik.",
    "Penggunaan sistem informasi dapat meningkatkan kualitas pengelolaan data di sekolah. Data yang tersimpan secara terstruktur lebih mudah dicari, diperbarui, dan direkap sesuai kebutuhan. Sistem membantu mengurangi kesalahan pencatatan karena setiap data dimasukkan melalui form yang memiliki aturan tertentu. Informasi yang tersedia pada halaman laporan juga membantu sekolah mengambil keputusan berdasarkan data yang lebih jelas.",
    "Berdasarkan kondisi tersebut, diperlukan perancangan Sistem Informasi Absensi PKL yang sesuai dengan kebutuhan sekolah. Sistem ini diharapkan membantu pencatatan kehadiran, monitoring oleh admin atau kesiswaan, pengelolaan data siswa dan lokasi PKL, serta pembuatan laporan. Perancangan dilakukan dengan pendekatan terstruktur agar alur data, proses, dan rancangan basis data dapat dijelaskan secara sistematis. Dengan adanya sistem ini, administrasi absensi PKL diharapkan menjadi lebih efektif, akurat, dan mudah digunakan.",
]
for x in latars:
    p(x)
h("1.2 Rumusan Masalah", 2)
p("Berdasarkan latar belakang yang telah dijelaskan, rumusan masalah disusun untuk menggambarkan kebutuhan utama pada proses absensi PKL. Rumusan masalah berfokus pada pencatatan kehadiran, monitoring siswa, pengelolaan data, dan penyajian laporan. Setiap rumusan diarahkan agar dapat dijawab melalui tujuan dan rancangan sistem. Rumusan masalah dalam laporan ini adalah sebagai berikut:")
numbered([
    "Bagaimana merancang Sistem Informasi Absensi PKL yang dapat membantu proses pencatatan kehadiran siswa secara lebih efektif?",
    "Bagaimana sistem dapat membantu admin atau kesiswaan dalam memantau kehadiran siswa selama menjalani PKL?",
    "Bagaimana sistem dapat mengelola data siswa, data lokasi PKL, dan data absensi secara terstruktur?",
    "Bagaimana sistem dapat menghasilkan laporan absensi PKL yang cepat, akurat, dan mudah dipahami?",
])
h("1.3 Tujuan", 2)
p("Tujuan penyusunan laporan ini disusun berdasarkan rumusan masalah yang telah ditentukan. Setiap tujuan diarahkan untuk menjawab kebutuhan sistem yang muncul dari proses pencatatan dan monitoring absensi PKL. Tujuan ini menjadi dasar dalam menentukan fitur utama yang perlu disediakan oleh sistem. Tujuan perancangan Sistem Informasi Absensi PKL adalah sebagai berikut:")
numbered([
    "Merancang Sistem Informasi Absensi PKL yang dapat membantu proses pencatatan kehadiran siswa secara efektif.",
    "Membantu admin atau kesiswaan dalam memantau kehadiran siswa selama menjalani PKL.",
    "Mengelola data siswa, data lokasi PKL, dan data absensi secara lebih terstruktur.",
    "Menghasilkan laporan absensi PKL yang cepat, akurat, dan mudah dipahami.",
])
h("1.4 Manfaat", 2)
for x in [
    "Manfaat sistem bagi siswa adalah memberikan kemudahan dalam mencatat kehadiran selama kegiatan praktik berlangsung. Siswa dapat melakukan absensi masuk dan pulang melalui sistem tanpa bergantung sepenuhnya pada lembar presensi manual. Proses tersebut membantu siswa membiasakan diri dengan kedisiplinan dan tanggung jawab dalam melaporkan kehadiran. Riwayat yang tersimpan juga dapat menjadi bukti bahwa siswa telah mengikuti kegiatan PKL sesuai jadwal.",
    "Bagi admin atau kesiswaan, sistem ini bermanfaat untuk memantau data kehadiran siswa secara lebih cepat dan terstruktur. Admin dapat melihat data siswa, lokasi PKL, status absensi, dan riwayat kehadiran melalui halaman yang tersedia. Kesiswaan juga dapat mengetahui siswa yang hadir, terlambat, izin, sakit, atau belum melakukan absensi. Dengan demikian, proses pengawasan dilakukan berdasarkan data yang tersimpan di dalam sistem.",
    "Bagi guru pembimbing, sistem menyediakan informasi kehadiran siswa yang dapat digunakan untuk proses pembinaan. Guru pembimbing dapat melihat kondisi kehadiran siswa tanpa menunggu laporan manual yang dikumpulkan secara terpisah. Informasi tersebut membantu guru mengambil tindakan apabila terdapat siswa yang sering terlambat atau tidak hadir. Sistem juga mendukung evaluasi karena data kehadiran dapat dijadikan dasar penilaian kedisiplinan.",
    "Bagi sekolah, sistem ini bermanfaat untuk meningkatkan kualitas administrasi pelaksanaan PKL. Data absensi yang tersimpan dalam database dapat digunakan untuk membuat laporan secara lebih cepat dan akurat. Sekolah dapat mengurangi risiko kehilangan data, keterlambatan rekap, dan kesalahan pencatatan yang sering terjadi pada proses manual. Dengan adanya sistem ini, pelaksanaan PKL dapat dikelola secara lebih modern, tertib, dan sesuai kebutuhan sekolah.",
]:
    p(x)
new_section()

# BAB II
h("BAB II\nKAJIAN PUSTAKA", 1)
topics = [
    ("2.1 Pengertian Sistem Informasi", "sistem informasi", "Novianti, Lestari, dan Raharjo (2022)"),
    ("2.2 Pengertian Absensi", "absensi", "Nadeak, Hutauruk, dan Prayudani (2023)"),
    ("2.3 Pengertian Praktik Kerja Lapangan (PKL)", "Praktik Kerja Lapangan", "Syuhada dan Muhathir (2024)"),
    ("2.4 Sistem Informasi Absensi PKL", "sistem informasi absensi PKL", "Nadeak et al. (2023)"),
    ("2.5 Pengertian Perancangan Sistem", "perancangan sistem", "Falipurnawati dan Darmadi (2019)"),
    ("2.6 Metode Perancangan Terstruktur", "metode perancangan terstruktur", "Falipurnawati dan Darmadi (2019)"),
    ("2.7 Flowchart", "flowchart", "Falipurnawati dan Darmadi (2019)"),
    ("2.8 Data Flow Diagram (DFD)", "Data Flow Diagram", "Falipurnawati dan Darmadi (2019)"),
    ("2.9 Entity Relationship Diagram (ERD)", "Entity Relationship Diagram", "Novianti et al. (2022)"),
    ("2.10 Basis Data", "basis data", "Armando dan Hermansyah (2025)"),
    ("2.11 Black Box Testing", "black box testing", "Muzakki, Nuryasin, dan Sumadi (2024)"),
]
for title, concept, cite in topics:
    h(title, 2)
    p(f"{concept.capitalize()} merupakan bagian penting dalam pengembangan Sistem Informasi Absensi PKL karena berhubungan dengan kebutuhan data dan proses yang digunakan oleh sekolah. {cite} menjelaskan bahwa penerapan sistem berbasis teknologi dapat membantu pekerjaan yang sebelumnya dilakukan secara manual menjadi lebih teratur. Pemahaman terhadap {concept} diperlukan agar rancangan sistem tidak hanya menampilkan fitur, tetapi juga menyelesaikan masalah pengguna. Dengan dasar teori yang tepat, sistem dapat dirancang sesuai kebutuhan pencatatan, monitoring, dan pelaporan absensi PKL.")
    p(f"Dalam laporan ini, {concept} digunakan sebagai landasan untuk menjelaskan bagaimana Sistem Informasi Absensi PKL bekerja. Konsep tersebut membantu menjelaskan proses yang terjadi mulai dari pengguna memasukkan data sampai sistem menghasilkan informasi. Data yang dikelola meliputi data siswa, lokasi PKL, absensi, validasi, dan laporan. Dengan memahami {concept}, rancangan sistem menjadi lebih mudah dibaca dan dapat dipertanggungjawabkan secara akademik.")
new_section()

# BAB III
h("BAB III\nANALISIS DAN PERANCANGAN", 1)
h("3.1 Analisis Sistem Berjalan", 2)
p("Sistem berjalan pada proses absensi PKL masih dapat dilakukan secara manual melalui pencatatan terpisah di tempat praktik atau laporan yang dikumpulkan kepada sekolah. Kondisi tersebut membuat data kehadiran tidak selalu dapat diketahui secara cepat oleh admin, kesiswaan, dan guru pembimbing. Proses rekap membutuhkan waktu karena data harus diperiksa dari banyak sumber sebelum dijadikan laporan. Oleh karena itu, sistem berjalan perlu dianalisis agar kelemahannya dapat dijadikan dasar dalam menyusun sistem usulan.")
h("3.2 Analisis Sistem Usulan", 2)
p("Sistem usulan berupa Sistem Informasi Absensi PKL yang mengelola data siswa, lokasi PKL, absensi, validasi, dan laporan dalam satu aplikasi. Pengguna masuk melalui halaman login, kemudian sistem menampilkan menu sesuai role masing-masing. Siswa dapat melakukan absensi masuk dan pulang, sedangkan admin atau kesiswaan dapat mengelola data dan memantau laporan. Sistem usulan ini diharapkan mempercepat pencatatan, memudahkan monitoring, dan menghasilkan laporan yang lebih akurat.")
design_order = [
    ("DFD-Level-0-Konteks-AbsensiPKL-flowchart.drawio.png", "3.3 Flowchart Sistem", "Gambar 3. 1 Flowchart Sistem Absensi PKL"),
    ("DFD-Level-0-Konteks-AbsensiPKL-DFD Level 0 - Konteks.drawio.png", "3.4 Diagram Konteks", "Gambar 3. 2 Diagram Konteks"),
    ("DFD-Level-0-Konteks-AbsensiPKL-dfd level 1.drawio.png", "3.5 Data Flow Diagram Level 1", "Gambar 3. 3 Data Flow Diagram Level 1"),
    ("DFD-Level-0-Konteks-AbsensiPKL-1.0.drawio.png", "3.6 DFD Level 2 Proses Autentikasi", "Gambar 3. 4 DFD Level 2 Proses Autentikasi"),
    ("DFD-Level-0-Konteks-AbsensiPKL-2.0.png", "3.7 DFD Level 2 Proses Absensi", "Gambar 3. 5 DFD Level 2 Proses Absensi"),
    ("DFD-Level-0-Konteks-AbsensiPKL-3.0.drawio.png", "3.8 DFD Level 2 Proses Laporan Harian", "Gambar 3. 6 DFD Level 2 Proses Laporan Harian"),
    ("DFD-Level-0-Konteks-AbsensiPKL-4.0.drawio.png", "3.9 DFD Level 2 Proses Validasi", "Gambar 3. 7 DFD Level 2 Proses Validasi"),
    ("DFD-Level-0-Konteks-AbsensiPKL-5.0.drawio.png", "3.10 DFD Level 2 Proses Rekap Absensi", "Gambar 3. 8 DFD Level 2 Proses Rekap Absensi"),
    ("DFD-Level-0-Konteks-AbsensiPKL-6.0.drawio.png", "3.11 DFD Level 2 Proses Monitoring", "Gambar 3. 9 DFD Level 2 Proses Monitoring"),
    ("DFD-Level-0-Konteks-AbsensiPKL-7.0.drawio.png", "3.12 DFD Level 2 Proses Pengelolaan Data", "Gambar 3. 10 DFD Level 2 Proses Pengelolaan Data"),
    ("DFD-Level-0-Konteks-AbsensiPKL-erd.drawio.png", "3.13 Entity Relationship Diagram", "Gambar 3. 11 Entity Relationship Diagram"),
]
for file, sub, cap in design_order:
    path = DESIGN / file
    if path.exists():
        h(sub, 2)
        img(path, cap, explain(cap, "perancangan sistem"), 5.8)
new_section()

# BAB IV
h("BAB IV\nHASIL IMPLEMENTASI", 1)
h("4.1 Peran dan Tanggung Jawab Pengguna Sistem", 2)
p("Sistem Informasi Absensi PKL memiliki beberapa role pengguna yang disesuaikan dengan kebutuhan pengelolaan data. Siswa berperan melakukan absensi, melihat riwayat, dan mengisi informasi kegiatan PKL. Admin dan kesiswaan memiliki fungsi yang sama atau mirip dalam mengelola data siswa, lokasi PKL, monitoring kehadiran, dan laporan absensi. Guru pembimbing atau pihak terkait menggunakan informasi yang tersedia untuk membantu pengawasan dan evaluasi pelaksanaan PKL.")
screen_caps = [
    "Halaman Login", "Halaman Dashboard Superadmin", "Halaman Manajemen Pengguna", "Halaman Hak Akses Menu",
    "Halaman Backup dan Restore", "Halaman Notifikasi Sistem", "Halaman Dashboard Admin", "Halaman Data Siswa PKL",
    "Halaman Setting Website", "Halaman Profil Pengguna", "Halaman Lokasi PKL", "Halaman Riwayat Absensi",
    "Halaman Chatbot Asisten", "Halaman Absensi Masuk Siswa", "Halaman Absensi Pulang Siswa", "Halaman Pengajuan Izin Siswa",
    "Halaman Validasi Pengajuan", "Halaman Laporan Bimbingan", "Halaman Validasi Catatan Bimbingan", "Halaman Daftar Pengajuan",
    "Halaman Summary Report Mingguan", "Halaman Analisis Mingguan", "Halaman Dashboard Pembimbing", "Halaman Summary Report Lanjutan",
    "Halaman Monitoring Progres", "Halaman Validasi Absensi", "Halaman Monitoring Siswa", "Halaman Riwayat Wali Kelas",
    "Halaman Laporan Analisis", "Halaman Dashboard Wali Kelas", "Halaman Dashboard Kepala Sekolah", "Halaman Laporan Grafik",
    "Halaman Dashboard Kesiswaan", "Halaman Validasi Kehadiran",
]
for idx, (path, title) in enumerate(zip(sorted(SCREENS.glob("*.png"), key=lambda x: x.name), screen_caps), 1):
    h(f"4.{idx+1} {title}", 2)
    img(path, f"Gambar 4. {idx} {title}", explain(title, "halaman aplikasi"), 5.8)
h("4.36 Hasil Black Box Testing", 2)
p("Pengujian black box testing dilakukan untuk memastikan setiap fungsi utama Sistem Informasi Absensi PKL berjalan sesuai kebutuhan. Pengujian ini dilakukan berdasarkan tampilan dan fungsi sistem tanpa melihat kode program. Setiap fitur diuji dengan memberikan input tertentu dan membandingkan output sistem dengan hasil yang diharapkan. Tabel berikut menunjukkan hasil pengujian terhadap beberapa fitur utama yang tersedia pada aplikasi.")
table(["No","Fitur yang Diuji","Skenario Pengujian","Input","Hasil yang Diharapkan","Hasil Pengujian","Status"], [
    [1,"Login pengguna","Pengguna memasukkan akun valid.","Username dan password benar.","Sistem menampilkan dashboard sesuai role.","Dashboard tampil sesuai akun.","Berhasil"],
    [2,"Menampilkan dashboard","Pengguna membuka dashboard.","Akun aktif.","Ringkasan data dan menu tampil.","Ringkasan tampil sesuai hak akses.","Berhasil"],
    [3,"Mengelola data siswa PKL","Admin menambah atau mengubah siswa.","Data identitas siswa.","Data tersimpan pada tabel.","Data berhasil tersimpan.","Berhasil"],
    [4,"Mengelola data tempat PKL","Admin memasukkan lokasi PKL.","Nama, alamat, dan lokasi.","Data lokasi tersimpan.","Lokasi tampil pada tabel.","Berhasil"],
    [5,"Melakukan absensi masuk","Siswa menekan tombol check in.","Koordinat dan waktu.","Absensi masuk tersimpan.","Riwayat absensi bertambah.","Berhasil"],
    [6,"Melakukan absensi pulang","Siswa menekan tombol check out.","Koordinat dan laporan kegiatan.","Absensi pulang tersimpan.","Data harian menjadi lengkap.","Berhasil"],
    [7,"Melihat riwayat absensi","Pengguna membuka riwayat.","Filter tanggal.","Daftar absensi tampil.","Riwayat tampil dalam tabel.","Berhasil"],
    [8,"Menampilkan laporan absensi","Pengguna memilih periode laporan.","Periode laporan.","Rekap absensi tampil.","Laporan dan grafik tampil.","Berhasil"],
    [9,"Mengelola hak akses","Superadmin mengatur checklist menu.","Pilihan role dan menu.","Hak akses tersimpan.","Menu berubah sesuai pengaturan.","Berhasil"],
    [10,"Logout","Pengguna keluar dari sistem.","Tombol logout.","Sesi berakhir dan kembali ke login.","Pengguna keluar dari sistem.","Berhasil"],
])
p("Berdasarkan hasil pengujian, fitur utama pada Sistem Informasi Absensi PKL berjalan sesuai skenario yang telah ditentukan. Pengujian dilakukan dengan melihat kesesuaian antara input yang diberikan dan output yang ditampilkan oleh sistem. Setiap fitur yang diuji menunjukkan hasil yang sesuai dengan kebutuhan, mulai dari login, pengelolaan data, absensi, laporan, hingga logout. Dengan demikian, sistem dapat dinyatakan layak digunakan secara fungsional untuk membantu proses pencatatan dan monitoring absensi PKL.")
new_section()

# BAB V
h("BAB V\nKESIMPULAN DAN SARAN", 1)
h("5.1 Kesimpulan", 2)
for x in [
    "Sistem Informasi Absensi PKL dapat membantu proses pencatatan kehadiran siswa secara lebih efektif. Sistem menyediakan fitur absensi masuk dan absensi pulang yang menggantikan pencatatan manual. Data yang dikirim oleh siswa dapat langsung tersimpan sehingga lebih mudah diperiksa kembali. Dengan demikian, tujuan pertama mengenai pencatatan kehadiran yang lebih efektif dapat tercapai melalui rancangan dan implementasi sistem.",
    "Sistem dapat membantu admin atau kesiswaan dalam memantau kehadiran siswa selama menjalani PKL. Halaman dashboard, riwayat absensi, validasi, dan laporan memberikan informasi yang dibutuhkan untuk melihat kondisi kehadiran siswa. Admin dan kesiswaan dapat mengetahui data yang masuk tanpa harus menunggu rekap manual dari tempat praktik. Dengan adanya fitur tersebut, monitoring kehadiran menjadi lebih cepat, terarah, dan mudah dilakukan.",
    "Pengelolaan data siswa, data lokasi PKL, dan data absensi dapat dilakukan secara lebih terstruktur melalui sistem. Data utama disimpan dalam database sehingga hubungan antar data lebih mudah dikelola. Admin dapat memperbarui data siswa dan lokasi PKL sesuai kebutuhan sekolah. Struktur data yang jelas membantu sistem menghasilkan informasi yang konsisten dan mengurangi risiko kesalahan pencatatan.",
    "Sistem mampu menghasilkan laporan absensi PKL yang cepat, akurat, dan mudah dipahami. Laporan dapat ditampilkan melalui tabel dan grafik sehingga pengguna dapat membaca informasi dengan lebih jelas. Data laporan berasal dari absensi yang tersimpan, sehingga penyusunan laporan tidak perlu dilakukan secara manual dari awal. Dengan demikian, sistem mendukung kebutuhan evaluasi dan administrasi pelaksanaan PKL di sekolah.",
]:
    p(x)
h("5.2 Saran", 2)
for x in [
    "Pengembangan sistem ke depan dapat dilakukan dengan menambahkan fitur notifikasi otomatis kepada siswa, admin, kesiswaan, dan guru pembimbing. Notifikasi dapat digunakan untuk mengingatkan siswa yang belum melakukan absensi masuk atau absensi pulang. Fitur ini juga dapat membantu kesiswaan mengetahui data yang perlu segera diperiksa. Dengan adanya notifikasi, proses monitoring dapat dilakukan secara lebih cepat dan responsif.",
    "Sistem juga dapat dikembangkan dengan integrasi lokasi GPS yang lebih rinci agar validasi kehadiran menjadi lebih kuat. Integrasi tersebut dapat membantu memastikan siswa melakukan absensi dari lokasi PKL yang sesuai. Selain itu, sistem dapat ditambahkan fitur export laporan dalam format PDF atau Excel agar laporan lebih mudah dibagikan. Pengembangan ini akan membuat sistem semakin sesuai dengan kebutuhan administrasi sekolah.",
    "Peningkatan keamanan akun perlu menjadi perhatian dalam pengembangan berikutnya. Sistem dapat menambahkan verifikasi dua langkah, pembatasan percobaan login, dan pencatatan aktivitas pengguna yang lebih lengkap. Keamanan penting karena sistem menyimpan data siswa, lokasi PKL, dan riwayat kehadiran yang perlu dilindungi. Dengan keamanan yang lebih baik, kepercayaan pengguna terhadap sistem juga akan meningkat.",
    "Selain berbasis web, sistem dapat dikembangkan menjadi aplikasi mobile agar lebih mudah digunakan oleh siswa saat berada di tempat PKL. Aplikasi mobile dapat mempermudah proses absensi karena siswa dapat mengakses fitur melalui perangkat yang sering digunakan. Pengembangan mobile juga dapat mendukung penggunaan kamera, lokasi, dan notifikasi perangkat secara lebih optimal. Dengan pengembangan tersebut, Sistem Informasi Absensi PKL dapat menjadi lebih praktis dan sesuai kebutuhan pengguna di lapangan.",
]:
    p(x)
new_section()

h("DAFTAR PUSTAKA", 1)
for ref in [
    "Armando, E., & Hermansyah, H. (2025). Rancang Bangun Sistem Absensi Online Karyawan dengan Validasi Foto sebagai Bukti Kehadiran. Digital Transformation Technology, 5(1), 250-257. https://doi.org/10.47709/digitech.v5i1.6025",
    "Falipurnawati, I., & Darmadi, E. A. (2019). Perancangan Proses Sistem Informasi Akademik Menggunakan Data Flow Diagram. In Search, 18(1). https://doi.org/10.37278/insearch.v18i1.145",
    "Muzakki, M. Z. I., Nuryasin, I., & Sumadi, F. D. S. (2024). Perancangan Sistem Informasi Absensi Menggunakan Smartcard Berbasis Internet of Things pada CV. Anugerah Mandiri. Jurnal Repositor. https://doi.org/10.22219/repositor.v3i4.32078",
    "Nadeak, Y. T., Hutauruk, M. Y., & Prayudani, S. (2023). Perancangan dan Pembuatan Absensi Berbasis Web di CV. SAE Digital Akademi. Prosiding Konferensi Nasional Social & Engineering Polmed, 4(1). https://doi.org/10.51510/konsep.v4i1.1252",
    "Novianti, Lestari, dan Raharjo. (2022). Sistem Informasi Pendataan dan Penilaian PKL Mahasiswa Berbasis Web. Journal Peqguruang: Conference Series. https://media.neliti.com/media/publications/359072-sistem-informasi-pendataan-dan-penilaian-48512eda.pdf",
    "Syuhada, R., & Muhathir. (2024). Pembuatan Sistem Absensi Siswa Praktik Kerja Lapangan Berbasis Web di CV Sae Akademi Digital Medan. INCODING: Journal of Informatics and Computer Science Engineering, 4(2). https://doi.org/10.34007/incoding.v4i2.729",
    "Wibawa, J. C., & Fathoni, M. (2021). Pengujian Black Box pada Aplikasi Sistem Absensi Karyawan Berbasis Website Menggunakan Metode Equivalent Partitions. Jurnal Informatika Universitas Pamulang. https://openjournal.unpam.ac.id/index.php/informatika/article/view/17696",
]:
    p(ref)

# setup_page_numbers()
if OUT.exists():
    OUT.unlink()
doc.save(OUT)
print(OUT)
